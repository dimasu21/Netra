"""
================================================================================
ASISTEN TINJAUAN PUSTAKA (Literature Review Assistant)
================================================================================
Aplikasi web untuk membantu Anda melakukan screening referensi fisik
(buku/jurnal) menggunakan OCR dan Algoritma Rabin-Karp.

Tech Stack:
- Backend: Flask (Python)
- OCR: Pytesseract (Tesseract OCR)
- Algorithm: Rabin-Karp dengan Rolling Hash

Author: Thesis Project
================================================================================
"""

import os
import re
import base64
from io import BytesIO
from flask import Flask, render_template, request, jsonify, session, redirect, url_for
from flask_login import LoginManager, current_user, login_required
from PIL import Image, ImageFile
ImageFile.LOAD_TRUNCATED_IMAGES = True
import pytesseract

# Document processing libraries
from docx import Document
import PyPDF2

# Groq AI with Fallback System (supports up to 5 API keys)
from groq import Groq
from dotenv import load_dotenv

# Database and Auth
from models import db, User, History
from datetime import datetime

# Load environment variables
load_dotenv()

# Configure Multiple Groq API Keys for Fallback
# Add your API keys to .env file: GROQ_API_KEY_1, GROQ_API_KEY_2, etc.
GROQ_API_KEYS = []
for i in range(1, 6):  # Support up to 5 API keys
    key = os.environ.get(f"GROQ_API_KEY_{i}")
    if key:
        GROQ_API_KEYS.append(key)

# Fallback to original single key if no numbered keys found
if not GROQ_API_KEYS:
    single_key = os.environ.get("GROQ_API_KEY")
    if single_key:
        GROQ_API_KEYS.append(single_key)

# Track current API key index
current_api_index = 0

def get_groq_client():
    """Get Groq client with current API key"""
    global current_api_index
    if not GROQ_API_KEYS:
        return None
    return Groq(api_key=GROQ_API_KEYS[current_api_index % len(GROQ_API_KEYS)])

def switch_to_next_api():
    """Switch to next available API key"""
    global current_api_index
    current_api_index = (current_api_index + 1) % len(GROQ_API_KEYS)
    print(f"[FALLBACK] Switching to API key #{current_api_index + 1}")

# Configure Tesseract path (auto-detect for Docker/Windows)
import shutil
tesseract_path = shutil.which('tesseract')
if tesseract_path:
    pytesseract.pytesseract.tesseract_cmd = tesseract_path
elif os.path.exists(r'C:\Program Files\Tesseract-OCR\tesseract.exe'):
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Fix for running behind reverse proxy (Cloudflare Tunnel)
from werkzeug.middleware.proxy_fix import ProxyFix

# Security & Limits
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_wtf.csrf import CSRFProtect

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'netra-secret-key-change-in-production')

# Database configuration - smart detection for local vs Docker
basedir = os.path.abspath(os.path.dirname(__file__))
local_db_path = 'sqlite:///' + os.path.join(basedir, 'instance', 'netra.db')

# Check if env DATABASE_URI is set and valid (for Docker deployment)
env_db_uri = os.environ.get('DATABASE_URI')
if env_db_uri and env_db_uri.startswith('sqlite:///'):
    # Extract path from URI and check if parent directory exists
    db_file_path = env_db_uri.replace('sqlite:///', '')
    # Only use env path if the parent directory actually exists on this system
    if os.path.exists(os.path.dirname(db_file_path)):
        app.config['SQLALCHEMY_DATABASE_URI'] = env_db_uri
    else:
        # Fallback to local path for development
        app.config['SQLALCHEMY_DATABASE_URI'] = local_db_path
elif env_db_uri:
    # Non-SQLite database (e.g., PostgreSQL) - use as is
    app.config['SQLALCHEMY_DATABASE_URI'] = env_db_uri
else:
    app.config['SQLALCHEMY_DATABASE_URI'] = local_db_path

app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['PREFERRED_URL_SCHEME'] = 'https'  # Force HTTPS for OAuth redirect URIs

# Initialize Security
csrf = CSRFProtect(app)
limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=["200 per day", "50 per hour"],
    storage_uri="memory://"
)
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_prefix=1)

# Initialize extensions
db.init_app(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'auth.login'
login_manager.login_message = 'Silakan login untuk mengakses halaman ini.'
login_manager.login_message_category = 'info'

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Register auth blueprint
from auth import auth as auth_blueprint, init_google_oauth
app.register_blueprint(auth_blueprint)

# Create database tables
with app.app_context():
    db.create_all()

# Initialize Google OAuth
init_google_oauth(app)


# ==============================================================================
# RABIN-KARP ALGORITHM IMPLEMENTATION
# ==============================================================================
"""
ALGORITMA RABIN-KARP dengan ROLLING HASH
=========================================

Rabin-Karp adalah algoritma pencarian string yang menggunakan fungsi hash
untuk menemukan pola dalam teks. Keunggulannya adalah efisiensi O(n+m)
dalam kasus rata-rata.

RUMUS ROLLING HASH:
-------------------
Untuk string s[0..m-1], nilai hash dihitung sebagai:

    H(s) = (s[0] × BASE^(m-1) + s[1] × BASE^(m-2) + ... + s[m-1] × BASE^0) mod PRIME

Dimana:
    - BASE = 256 (jumlah karakter ASCII)
    - PRIME = 101 (bilangan prima untuk mengurangi collision)
    - m = panjang pattern yang dicari

ROLLING UPDATE:
---------------
Saat menggeser window dari posisi i ke i+1:

    H_new = ((H_old - s[i] × BASE^(m-1)) × BASE + s[i+m]) mod PRIME

Ini memungkinkan perhitungan hash O(1) untuk setiap posisi baru,
bukan O(m) jika dihitung ulang dari awal.

UNTUK SIDANG THESIS:
--------------------
1. Kompleksitas waktu rata-rata: O(n + m)
2. Kompleksitas waktu terburuk: O(nm) - saat banyak collision
3. Digunakan untuk: plagiarism detection, DNA sequence matching, dll.
"""

# Konstanta untuk Rabin-Karp
BASE = 256      # Jumlah karakter dalam alfabet ASCII
PRIME = 101     # Bilangan prima untuk operasi modulo


# ==============================================================================
# AI SUMMARY FUNCTION (GROQ)
# ==============================================================================

def generate_ai_summary(text: str) -> dict:
    """
    Generate AI analysis of journal/academic document quality using Groq.
    Returns structured JSON for tabs display.
    
    Args:
        text (str): Extracted text from document
        
    Returns:
        dict: Structured journal quality analysis in JSON format
    """
    try:
        prompt = f"""Anda adalah analis kualitas jurnal akademik profesional.

Analisis dokumen berikut secara objektif dan teknis.
JANGAN menulis ulang isi dokumen.
JANGAN meringkas dokumen.

Gunakan Bahasa Indonesia formal dan akademik.

TEKS DOKUMEN:
{text[:4500]}

BERIKAN OUTPUT DALAM FORMAT JSON YANG VALID (tanpa markdown code block):

{{
    "deteksi_struktur": {{
        "abstrak": "ada/tidak ada",
        "pendahuluan": "ada/tidak ada",
        "metodologi": "ada/tidak ada",
        "hasil_pembahasan": "ada/tidak ada",
        "kesimpulan": "ada/tidak ada"
    }},
    "kekuatan_metodologi": {{
        "skor": "Kuat/Sedang/Lemah",
        "penjelasan": "jelaskan dalam 2-3 kalimat mengapa metodologi kuat/sedang/lemah",
        "poin_positif": ["poin 1", "poin 2"]
    }},
    "kejelasan_variabel": {{
        "status": "Jelas/Kurang Jelas/Tidak Jelas",
        "variabel_teridentifikasi": ["variabel 1", "variabel 2"],
        "catatan": "jelaskan jika ada masalah kejelasan variabel"
    }},
    "konsistensi": {{
        "tujuan_vs_metodologi": "Sesuai/Kurang Sesuai/Tidak Sesuai",
        "metodologi_vs_hasil": "Sesuai/Kurang Sesuai/Tidak Sesuai",
        "tujuan_vs_kesimpulan": "Sesuai/Kurang Sesuai/Tidak Sesuai",
        "catatan": "jelaskan inkonsistensi jika ada"
    }},
    "kelemahan": {{
        "struktural": ["kelemahan 1", "kelemahan 2"],
        "metodologis": ["kelemahan 1", "kelemahan 2"],
        "penulisan": ["kelemahan 1", "kelemahan 2"]
    }},
    "skor_keseluruhan": {{
        "nilai": "A/B/C/D",
        "interpretasi": "Sangat Baik/Baik/Cukup/Perlu Perbaikan",
        "rekomendasi": "saran singkat untuk perbaikan"
    }}
}}

PENTING: Output harus berupa JSON valid tanpa markdown code block, tanpa backticks, langsung JSON saja."""
        
        # Use Groq API with Fallback System
        max_retries = len(GROQ_API_KEYS) if GROQ_API_KEYS else 1
        last_error = None
        
        for attempt in range(max_retries):
            try:
                client = get_groq_client()
                if not client:
                    raise Exception("No API keys configured")
                
                response = client.chat.completions.create(
                    model="llama-3.1-8b-instant",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.2,
                    max_tokens=1000,
                    response_format={"type": "json_object"}
                )
                
                ai_text = response.choices[0].message.content
                break  # Success, exit retry loop
                
            except Exception as e:
                last_error = e
                error_str = str(e).lower()
                
                # Check if rate limit error (429)
                if "rate_limit" in error_str or "429" in error_str or "quota" in error_str:
                    print(f"[RATE LIMIT] API key #{current_api_index + 1} hit limit, trying next...")
                    switch_to_next_api()
                    continue
                else:
                    # Other error, don't retry
                    raise e
        else:
            # All APIs exhausted
            raise Exception(f"All {max_retries} API keys exhausted. Last error: {last_error}")
        
        # Try to parse JSON, fallback to raw text if fails
        import json
        try:
            # Clean potential markdown code blocks
            clean_text = ai_text.strip()
            if clean_text.startswith('```'):
                clean_text = clean_text.split('```')[1]
                if clean_text.startswith('json'):
                    clean_text = clean_text[4:]
            if clean_text.endswith('```'):
                clean_text = clean_text[:-3]
            
            parsed_json = json.loads(clean_text.strip())
            return {
                'success': True,
                'summary': ai_text,
                'parsed': parsed_json,
                'is_json': True
            }
        except json.JSONDecodeError:
            return {
                'success': True,
                'summary': ai_text,
                'parsed': None,
                'is_json': False
            }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'summary': None,
            'parsed': None,
            'is_json': False
        }


def rabin_karp_search(text: str, pattern: str) -> list:
    """
    Implementasi MANUAL Algoritma Rabin-Karp untuk mencari pattern dalam teks.
    TIDAK menggunakan str.find() atau str.index().
    
    Args:
        text (str): Teks yang akan dicari
        pattern (str): Pattern/keyword yang ingin ditemukan
        
    Returns:
        list: Daftar posisi (index) dimana pattern ditemukan
    """
    # Handle edge cases
    if not pattern or not text:
        return []
    
    if len(pattern) > len(text):
        return []
    
    # Case-insensitive search: konversi ke lowercase
    text_lower = text.lower()
    pattern_lower = pattern.lower()
    
    n = len(text_lower)      # Panjang teks
    m = len(pattern_lower)   # Panjang pattern
    
    positions = []  # List untuk menyimpan posisi yang ditemukan
    
    # =========================================================================
    # LANGKAH 1: Hitung nilai h = BASE^(m-1) mod PRIME
    # Nilai ini digunakan untuk menghapus digit terdepan saat rolling
    # =========================================================================
    h = 1
    for i in range(m - 1):
        h = (h * BASE) % PRIME
    
    # =========================================================================
    # LANGKAH 2: Hitung hash awal untuk pattern dan window pertama teks
    # Menggunakan rumus: H = (c[0]*BASE^(m-1) + c[1]*BASE^(m-2) + ... + c[m-1])
    # =========================================================================
    pattern_hash = 0
    text_hash = 0
    
    for i in range(m):
        # ord() mengkonversi karakter ke nilai ASCII
        pattern_hash = (BASE * pattern_hash + ord(pattern_lower[i])) % PRIME
        text_hash = (BASE * text_hash + ord(text_lower[i])) % PRIME
    
    # =========================================================================
    # LANGKAH 3: Slide pattern over text dan bandingkan hash
    # =========================================================================
    for i in range(n - m + 1):
        # Jika hash sama, verifikasi dengan perbandingan karakter
        # (untuk menghindari false positive akibat collision)
        if pattern_hash == text_hash:
            # Verifikasi character-by-character
            match = True
            for j in range(m):
                if text_lower[i + j] != pattern_lower[j]:
                    match = False
                    break
            
            if match:
                positions.append(i)
        
        # =====================================================================
        # LANGKAH 4: Hitung hash untuk window berikutnya (Rolling Hash)
        # Rumus: H_new = ((H_old - leading_char * h) * BASE + trailing_char) mod PRIME
        # =====================================================================
        if i < n - m:
            # Hapus karakter terdepan, geser, dan tambah karakter baru
            leading_char = ord(text_lower[i])
            trailing_char = ord(text_lower[i + m])
            
            text_hash = (BASE * (text_hash - leading_char * h) + trailing_char) % PRIME
            
            # Handle nilai hash negatif
            if text_hash < 0:
                text_hash += PRIME
    
    return positions


def highlight_text(text: str, pattern: str) -> str:
    """
    Highlight semua kemunculan pattern dalam teks dengan tag HTML <mark>.
    Menggunakan Rabin-Karp untuk menemukan posisi, bukan str.find().
    
    Args:
        text (str): Teks asli
        pattern (str): Pattern yang akan di-highlight
        
    Returns:
        str: Teks dengan pattern yang sudah di-highlight
    """
    positions = rabin_karp_search(text, pattern)
    
    if not positions:
        return text
    
    # Build hasil dengan mengganti dari belakang untuk menjaga index
    result = text
    pattern_len = len(pattern)
    
    for pos in sorted(positions, reverse=True):
        original = result[pos:pos + pattern_len]
        highlighted = f'<mark class="highlight">{original}</mark>'
        result = result[:pos] + highlighted + result[pos + pattern_len:]
    
    return result


# ==============================================================================
# OCR FUNCTION
# ==============================================================================

def extract_text_from_image(image_data) -> str:
    """
    Ekstrak teks dari gambar menggunakan Tesseract OCR.
    
    Args:
        image_data: Data gambar (file atau bytes)
        
    Returns:
        str: Teks yang diekstrak dari gambar
    """
    try:
        image = Image.open(image_data)
        
        # Konversi ke RGB jika diperlukan
        if image.mode in ('RGBA', 'P'):
            image = image.convert('RGB')
        
        # Ekstrak teks dengan Tesseract
        text = pytesseract.image_to_string(image, lang='eng+ind')
        
        return text.strip()
    
    except Exception as e:
        raise Exception(f"OCR Error: {str(e)}")


def extract_text_from_docx(file_data) -> str:
    """
    Ekstrak teks dari file DOCX.
    
    Args:
        file_data: Data file DOCX (BytesIO)
        
    Returns:
        str: Teks yang diekstrak
    """
    try:
        doc = Document(file_data)
        full_text = []
        
        # Ekstrak dari paragraf
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Ekstrak dari tabel
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(' | '.join(row_text))
        
        return '\n'.join(full_text)
    
    except Exception as e:
        raise Exception(f"DOCX Error: {str(e)}")


def extract_text_from_pdf(file_data) -> str:
    """
    Ekstrak teks dari file PDF.
    
    Args:
        file_data: Data file PDF (BytesIO)
        
    Returns:
        str: Teks yang diekstrak
    """
    try:
        pdf_reader = PyPDF2.PdfReader(file_data)
        full_text = []
        
        for page_num, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text:
                full_text.append(f"--- Halaman {page_num + 1} ---\n{page_text}")
        
        return '\n\n'.join(full_text)
    
    except Exception as e:
        raise Exception(f"PDF Error: {str(e)}")


def get_file_type(filename: str) -> str:
    """Menentukan tipe file berdasarkan ekstensi."""
    if not filename:
        return 'unknown'
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    if ext in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff']:
        return 'image'
    elif ext in ['doc', 'docx']:
        return 'docx'
    elif ext == 'pdf':
        return 'pdf'
    return 'unknown'


# ==============================================================================
# FLASK ROUTES
# ==============================================================================

@app.route('/robots.txt')
def robots():
    return app.send_static_file('robots.txt')

@app.route('/sitemap.xml')
def sitemap():
    return app.send_static_file('sitemap.xml')

@app.route('/')
def index():
    """Render halaman utama."""
    return render_template('index.html')


@app.route('/pricing')
def pricing():
    """Render halaman pricing."""
    return render_template('pricing.html')


@app.route('/about')
def about():
    """Render halaman about."""
    return render_template('about.html')


@app.route('/contact')
def contact():
    """Render halaman contact."""
    return render_template('contact.html')


@app.route('/privacy')
def privacy():
    """Render halaman Kebijakan Privasi."""
    return render_template('privacy.html')


@app.route('/terms')
def terms():
    """Render halaman Syarat & Ketentuan."""
    return render_template('terms.html')


@app.route('/profile')
@login_required
def profile():
    """Render halaman profil user."""
    # Get recent activities from history
    activities = []
    if current_user.is_authenticated:
        from models import History
        recent_history = History.query.filter_by(user_id=current_user.id).order_by(History.created_at.desc()).limit(5).all()
        for h in recent_history:
            activities.append({
                'action': h.action_type or 'Analisis',
                'filename': h.filename or 'document',
                'time': h.created_at.strftime('%d %b %Y, %H:%M') if h.created_at else 'baru saja'
            })
    return render_template('profile.html', activities=activities)


@app.route('/batch')
def batch():
    """Render halaman batch processing."""
    return render_template('batch.html')


@app.route('/ocr')
def ocr():
    """Render halaman OCR Scanner."""
    return render_template('ocr.html')


@app.route('/document')
def document():
    """Render halaman Document Parser."""
    return render_template('document.html')


@app.route('/history')
def history():
    """Render halaman History/Riwayat Analisis."""
    if not current_user.is_authenticated:
        return redirect(url_for('auth.login'))
    
    # Ambil data history user, urutkan dari yang terbaru
    histories = History.query.filter_by(user_id=current_user.id).order_by(History.created_at.desc()).all()
    return render_template('history.html', histories=histories)
    
@app.route('/history/<int:id>/delete', methods=['POST'])
def delete_history(id):
    if not current_user.is_authenticated:
        return jsonify({'success': False}), 403
        
    history = History.query.get_or_404(id)
    if history.user_id != current_user.id:
        return jsonify({'success': False}), 403
        
    db.session.delete(history)
    db.session.commit()
    return jsonify({'success': True})


# Login and Signup routes are now handled by auth blueprint


# ==============================================================================
# BATCH RATE LIMITING API
# ==============================================================================

@app.route('/api/batch-limit/check', methods=['GET'])
@csrf.exempt
def check_batch_limit():
    """Check remaining batch usage for current user."""
    if not current_user.is_authenticated:
        # Guest users have unlimited batch (for now, could limit by IP later)
        return jsonify({
            'allowed': True,
            'remaining': 999,
            'limit': 999,
            'message': 'Guest user - unlimited'
        })
    
    from models import UsageLimit
    usage = UsageLimit.get_or_create(current_user.id, 'batch')
    
    # Check limit without incrementing
    from datetime import datetime, timedelta
    now = datetime.utcnow()
    
    # Reset if 24 hours passed
    if usage.last_reset is None or (now - usage.last_reset) > timedelta(hours=24):
        remaining = UsageLimit.BATCH_DAILY_LIMIT
        used = 0
    else:
        remaining = max(0, UsageLimit.BATCH_DAILY_LIMIT - usage.usage_count)
        used = usage.usage_count
    
    return jsonify({
        'allowed': remaining > 0,
        'remaining': remaining,
        'used': used,
        'limit': UsageLimit.BATCH_DAILY_LIMIT,
        'message': f'{remaining} batch tersisa hari ini' if remaining > 0 else 'Limit tercapai'
    })


@app.route('/api/batch-limit/use', methods=['POST'])
@csrf.exempt
def use_batch_limit():
    """Use one batch quota. Call this before processing batch."""
    if not current_user.is_authenticated:
        return jsonify({
            'allowed': True,
            'remaining': 999,
            'message': 'Guest user - unlimited'
        })
    
    from models import UsageLimit
    usage = UsageLimit.get_or_create(current_user.id, 'batch')
    allowed, remaining, message = usage.check_and_increment()
    
    return jsonify({
        'allowed': allowed,
        'remaining': remaining,
        'limit': UsageLimit.BATCH_DAILY_LIMIT,
        'message': message
    })


@app.route('/api/analyze', methods=['POST'])
@csrf.exempt
@limiter.limit("20 per minute")
def analyze():
    """
    API endpoint untuk menganalisis relevansi dokumen.
    Mendukung: Image (JPG, PNG), DOCX, PDF
    """
    try:
        # Validasi input - support both 'file' and 'image' keys
        uploaded_file = request.files.get('file') or request.files.get('image')
        
        if not uploaded_file:
            return jsonify({'success': False, 'error': 'Tidak ada file yang diunggah'}), 400
        
        # Cloudflare Turnstile Verification
        # Skip if user is logged in, already verified in session, OR running in development mode
        is_localhost = request.host.startswith('127.0.0.1') or request.host.startswith('localhost')
        
        if not current_user.is_authenticated and not session.get('is_human_verified') and not is_localhost:
            turnstile_token = request.form.get('cf-turnstile-response')
            if not turnstile_token:
                return jsonify({'success': False, 'error': 'Verifikasi keamanan gagal (Token missing). Silakan refresh halaman.'}), 400

            # Verify token with Cloudflare
            try:
                import requests
                secret_key = os.environ.get('TURNSTILE_SECRET_KEY', '0x4AAAAAACH9KxS6HutyBiLRk8STEbKs-j8')
                verify_response = requests.post(
                    'https://challenges.cloudflare.com/turnstile/v0/siteverify',
                    data={
                        'secret': secret_key,
                        'response': turnstile_token,
                        'remoteip': request.remote_addr
                    }
                ).json()

                if not verify_response.get('success'):
                    return jsonify({'success': False, 'error': 'Verifikasi keamanan gagal. Silakan coba lagi.'}), 400
                
                # Mark session as verified
                session['is_human_verified'] = True
                
            except Exception as e:
                # Fallback if verification api fails
                pass

        keyword = request.form.get('keyword', '').strip()
        
        if uploaded_file.filename == '':
            return jsonify({'success': False, 'error': 'Tidak ada file yang dipilih'}), 400
        
        if not keyword:
            return jsonify({'success': False, 'error': 'Variabel riset tidak boleh kosong'}), 400
        
        # Tentukan tipe file
        filename = uploaded_file.filename
        file_type = get_file_type(filename)
        file_bytes = BytesIO(uploaded_file.read())
        
        # Ekstrak teks berdasarkan tipe file
        if file_type == 'image':
            # Optimize image
            file_bytes.seek(0)
            img = Image.open(file_bytes)
            if img.mode == 'RGBA':
                img = img.convert('RGB')
            
            # Resize if too large
            if img.width > 2000 or img.height > 2000:
                img.thumbnail((2000, 2000), Image.Resampling.LANCZOS)
                # Save optimized to buffer
                buf = BytesIO()
                img.save(buf, format='JPEG', quality=85)
                file_bytes = BytesIO(buf.getvalue())
            
            file_bytes.seek(0)
            extracted_text = extract_text_from_image(file_bytes)
            # Generate preview untuk gambar
            file_bytes.seek(0)
            image_base64 = base64.b64encode(file_bytes.read()).decode('utf-8')
            file_bytes.seek(0)
            img = Image.open(file_bytes)
            mime = f"image/{img.format.lower()}" if img.format else "image/png"
            file_preview = f"data:{mime};base64,{image_base64}"
            
        elif file_type == 'docx':
            extracted_text = extract_text_from_docx(file_bytes)
            file_preview = None  # Tidak ada preview untuk dokumen
            
        elif file_type == 'pdf':
            extracted_text = extract_text_from_pdf(file_bytes)
            file_preview = None  # Tidak ada preview untuk PDF
            
        else:
            return jsonify({
                'success': False, 
                'error': 'Tipe file tidak didukung. Gunakan: JPG, PNG, DOCX, atau PDF'
            }), 400
        
        if not extracted_text:
            return jsonify({
                'success': True,
                'is_relevant': False,
                'extracted_text': '(Tidak ada teks yang terdeteksi)',
                'highlighted_text': '(Tidak ada teks yang terdeteksi)',
                'match_count': 0,
                'positions': [],
                'image_preview': None,
                'file_type': file_type,
                'filename': filename
            })
        
        # Cari keyword menggunakan Rabin-Karp
        positions = rabin_karp_search(extracted_text, keyword)
        is_relevant = len(positions) > 0
        
        # Highlight teks jika ditemukan
        highlighted_text = highlight_text(extracted_text, keyword) if is_relevant else extracted_text
        
        # Generate AI Summary
        ai_result = generate_ai_summary(extracted_text)
        
        # Save to History if User is Authenticated
        if current_user.is_authenticated:
            try:
                new_history = History(
                    user_id=current_user.id,
                    filename=filename,
                    file_type=file_type.upper(),
                    extracted_text=extracted_text,
                    ai_summary=ai_result.get('summary') if ai_result.get('success') else None,
                    created_at=datetime.utcnow()
                )
                db.session.add(new_history)
                db.session.commit()
            except Exception as e:
                print(f"Failed to save history: {e}")
        
        return jsonify({
            'success': True,
            'is_relevant': is_relevant,
            'extracted_text': extracted_text,
            'highlighted_text': highlighted_text,
            'match_count': len(positions),
            'positions': positions,
            'keyword': keyword,
            'image_preview': file_preview,
            'file_type': file_type,
            'filename': filename,
            'ai_summary': ai_result.get('summary') if ai_result.get('success') else None,
            'ai_parsed': ai_result.get('parsed') if ai_result.get('success') else None,
            'ai_is_json': ai_result.get('is_json', False),
            'ai_error': ai_result.get('error') if not ai_result.get('success') else None
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 400


# ==============================================================================
# ADMIN PANEL ROUTES
# ==============================================================================

def admin_required(f):
    """Decorator untuk membatasi akses hanya untuk admin."""
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated:
            flash('Silakan login terlebih dahulu.', 'warning')
            return redirect(url_for('auth.login'))
        if not current_user.is_admin:
            flash('Akses ditolak. Anda bukan administrator.', 'danger')
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function


@app.route('/admin')
@admin_required
def admin_dashboard():
    """Admin dashboard - lihat semua user."""
    users = User.query.order_by(User.created_at.desc()).all()
    return render_template('admin.html', users=users)


@app.route('/admin/user/<int:user_id>/toggle-active', methods=['POST'])
@admin_required
@csrf.exempt
def admin_toggle_user_active(user_id):
    """Toggle status aktif user."""
    user = User.query.get_or_404(user_id)
    if user.id == current_user.id:
        return jsonify({'success': False, 'error': 'Tidak bisa menonaktifkan diri sendiri'}), 400
    user.is_active = not user.is_active
    db.session.commit()
    return jsonify({'success': True, 'is_active': user.is_active})


@app.route('/admin/user/<int:user_id>/toggle-admin', methods=['POST'])
@admin_required
@csrf.exempt
def admin_toggle_user_admin(user_id):
    """Toggle status admin user."""
    user = User.query.get_or_404(user_id)
    if user.id == current_user.id:
        return jsonify({'success': False, 'error': 'Tidak bisa mengubah status admin diri sendiri'}), 400
    user.is_admin = not user.is_admin
    db.session.commit()
    return jsonify({'success': True, 'is_admin': user.is_admin})


@app.route('/admin/user/<int:user_id>/delete', methods=['POST'])
@admin_required
@csrf.exempt
def admin_delete_user(user_id):
    """Hapus user dari database."""
    user = User.query.get_or_404(user_id)
    if user.id == current_user.id:
        return jsonify({'success': False, 'error': 'Tidak bisa menghapus diri sendiri'}), 400
    # Hapus history user terlebih dahulu
    History.query.filter_by(user_id=user.id).delete()
    db.session.delete(user)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/admin/user/<int:user_id>/reset-password', methods=['POST'])
@admin_required
@csrf.exempt
def admin_reset_password(user_id):
    """Reset password user ke default."""
    user = User.query.get_or_404(user_id)
    if user.google_id:
        return jsonify({'success': False, 'error': 'User Google tidak bisa reset password'}), 400
    # Set password default
    new_password = 'netra123'
    user.set_password(new_password)
    db.session.commit()
    return jsonify({'success': True, 'new_password': new_password})


# ==============================================================================
# ERROR HANDLERS
# ==============================================================================

@app.errorhandler(404)
def page_not_found(e):
    return render_template('errors/404.html'), 404

@app.errorhandler(500)
def internal_server_error(e):
    return render_template('errors/500.html'), 500

# ==============================================================================
# MAIN ENTRY POINT
# ==============================================================================

if __name__ == '__main__':
    print("""
    ╔══════════════════════════════════════════════════════════════════════╗
    ║                    ASISTEN TINJAUAN PUSTAKA                          ║
    ║         Screening Referensi Fisik Berbasis OCR & Rabin-Karp          ║
    ╠══════════════════════════════════════════════════════════════════════╣
    ║   Teknologi:                                                         ║
    ║   • OCR: Tesseract (pytesseract)                                     ║
    ║   • Algorithm: Rabin-Karp dengan Rolling Hash                        ║
    ║   • Frontend: Bootstrap 5                                            ║
    ╠══════════════════════════════════════════════════════════════════════╣
    ║   Buka di browser: http://127.0.0.1:5000                             ║
    ╚══════════════════════════════════════════════════════════════════════╝
    """)
    app.run(debug=True, host='0.0.0.0', port=5000)
